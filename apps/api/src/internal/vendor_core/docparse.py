"""Document parsing, chunking, and content-hash deduplication.

A ``BaseParser`` ABC returning a uniform ``ParsedDocument``, concrete PDF/DOCX/HTML/
Markdown parsers (heavy deps dynamically imported, mirroring ``llm.py``), text
chunking strategies, and SHA-256 dedup helpers. This is the single largest
ingestion-side dedup win across the workspace.
"""

import hashlib
import re
from abc import ABC, abstractmethod
from enum import Enum
from typing import Any, Callable, Iterable, Optional, TypeVar, Union

from pydantic import BaseModel, Field

T = TypeVar("T")

_SENTENCE_RE = re.compile(r"(?<=[.!?])\s+")
_HEADING_SPLIT_RE = re.compile(r"\n(?=#{1,6}\s)")


class ParsedDocument(BaseModel):
    """Normalized output of parsing a document of any supported type."""

    text: str
    title: Optional[str] = None
    metadata: dict[str, Any] = Field(default_factory=dict)
    page_count: int = 0


class ParseResult(BaseModel):
    """Wraps a single parse attempt for batch flows (success or error)."""

    filename: Optional[str] = None
    document: Optional[ParsedDocument] = None
    error: Optional[str] = None

    @property
    def ok(self) -> bool:
        return self.document is not None


class BaseParser(ABC):
    """Abstract document parser."""

    supported_extensions: tuple[str, ...] = ()

    @abstractmethod
    def parse(
        self, data: bytes, *, filename: Optional[str] = None
    ) -> ParsedDocument: ...


class PdfParser(BaseParser):
    """PDF parsing via PyMuPDF (fitz)."""

    supported_extensions = (".pdf",)

    def parse(self, data: bytes, *, filename: Optional[str] = None) -> ParsedDocument:
        try:
            import fitz  # type: ignore  # PyMuPDF
        except ImportError:
            raise ImportError(
                "PyMuPDF is not installed. Install it via 'pip install pymupdf'."
            ) from None
        doc = fitz.open(stream=data, filetype="pdf")
        try:
            pages = [page.get_text() for page in doc]
            meta = doc.metadata or {}
            title = meta.get("title") or None
        finally:
            doc.close()
        return ParsedDocument(
            text="\n\n".join(pages),
            title=title,
            page_count=len(pages),
            metadata={"source": filename} if filename else {},
        )


class DocxParser(BaseParser):
    """DOCX parsing via python-docx."""

    supported_extensions = (".docx",)

    def parse(self, data: bytes, *, filename: Optional[str] = None) -> ParsedDocument:
        import io

        try:
            import docx  # type: ignore  # python-docx
        except ImportError:
            raise ImportError(
                "python-docx is not installed. "
                "Install it via 'pip install python-docx'."
            ) from None
        document = docx.Document(io.BytesIO(data))
        text = "\n".join(p.text for p in document.paragraphs)
        return ParsedDocument(
            text=text, metadata={"source": filename} if filename else {}
        )


class HtmlParser(BaseParser):
    """HTML parsing via BeautifulSoup."""

    supported_extensions = (".html", ".htm")

    def parse(self, data: bytes, *, filename: Optional[str] = None) -> ParsedDocument:
        try:
            from bs4 import BeautifulSoup  # type: ignore
        except ImportError:
            raise ImportError(
                "beautifulsoup4 is not installed. "
                "Install it via 'pip install beautifulsoup4'."
            ) from None
        soup = BeautifulSoup(data, "html.parser")
        title = soup.title.string if soup.title else None
        for tag in soup(["script", "style"]):
            tag.decompose()
        raw = soup.get_text(separator="\n")
        text = "\n".join(line.strip() for line in raw.splitlines() if line.strip())
        return ParsedDocument(
            text=text, title=title, metadata={"source": filename} if filename else {}
        )


class MarkdownParser(BaseParser):
    """Markdown / plain-text passthrough (no heavy dependency required)."""

    supported_extensions = (".md", ".markdown", ".txt")

    def parse(self, data: bytes, *, filename: Optional[str] = None) -> ParsedDocument:
        raw = (
            data.decode("utf-8", errors="replace")
            if isinstance(data, bytes)
            else str(data)
        )
        title = None
        for line in raw.splitlines():
            if line.startswith("# "):
                title = line[2:].strip()
                break
        return ParsedDocument(
            text=raw, title=title, metadata={"source": filename} if filename else {}
        )


_PARSERS: tuple[type[BaseParser], ...] = (
    PdfParser,
    DocxParser,
    HtmlParser,
    MarkdownParser,
)
_MIME_HINTS: dict[str, type[BaseParser]] = {
    "pdf": PdfParser,
    "word": DocxParser,
    "officedocument.wordprocessing": DocxParser,
    "html": HtmlParser,
    "markdown": MarkdownParser,
    "text/plain": MarkdownParser,
}


def get_parser(filename_or_mime: str) -> BaseParser:
    """Resolve a parser by file extension or MIME hint."""
    name = filename_or_mime.lower()
    for parser_cls in _PARSERS:
        if any(name.endswith(ext) for ext in parser_cls.supported_extensions):
            return parser_cls()
    for hint, parser_cls in _MIME_HINTS.items():
        if hint in name:
            return parser_cls()
    from .errors import ValidationError

    raise ValidationError(f"No parser available for '{filename_or_mime}'")


class ChunkStrategy(str, Enum):
    """Text chunking strategy."""

    FIXED = "fixed"
    SEMANTIC = "semantic"
    STRUCTURAL = "structural"


def chunk_text(
    text: str,
    *,
    strategy: ChunkStrategy = ChunkStrategy.FIXED,
    chunk_size: int = 512,
    overlap: int = 64,
) -> list[str]:
    """Split ``text`` into chunks using the requested strategy."""
    if not text:
        return []
    if strategy == ChunkStrategy.STRUCTURAL:
        return _chunk_structural(text, chunk_size)
    if strategy == ChunkStrategy.SEMANTIC:
        return _chunk_semantic(text, chunk_size)
    return _chunk_fixed(text, chunk_size, overlap)


def _chunk_fixed(text: str, chunk_size: int, overlap: int) -> list[str]:
    if overlap >= chunk_size:
        overlap = chunk_size // 4
    chunks: list[str] = []
    start = 0
    length = len(text)
    while start < length:
        end = min(start + chunk_size, length)
        chunks.append(text[start:end])
        if end >= length:
            break
        start = end - overlap
    return chunks


def _chunk_semantic(text: str, chunk_size: int) -> list[str]:
    sentences = _SENTENCE_RE.split(text)
    chunks: list[str] = []
    current = ""
    for sentence in sentences:
        if current and len(current) + len(sentence) + 1 > chunk_size:
            chunks.append(current.strip())
            current = sentence
        else:
            current = f"{current} {sentence}".strip()
    if current.strip():
        chunks.append(current.strip())
    return chunks


def _chunk_structural(text: str, chunk_size: int) -> list[str]:
    blocks = _HEADING_SPLIT_RE.split(text)
    chunks: list[str] = []
    for block in blocks:
        stripped = block.strip()
        if not stripped:
            continue
        if len(stripped) <= chunk_size:
            chunks.append(stripped)
        else:
            chunks.extend(_chunk_fixed(stripped, chunk_size, chunk_size // 8))
    return chunks


def compute_hash(content: Union[str, bytes]) -> str:
    """Return the SHA-256 hex digest of ``content``."""
    if isinstance(content, str):
        content = content.encode("utf-8")
    return hashlib.sha256(content).hexdigest()


def is_duplicate(content_hash: str, seen: set[str]) -> bool:
    """Return True if ``content_hash`` is already present in ``seen``."""
    return content_hash in seen


def filter_duplicates(items: Iterable[T], key: Callable[[T], str]) -> list[T]:
    """Return ``items`` with later duplicates (by ``key``) removed, preserving order."""
    seen: set[str] = set()
    result: list[T] = []
    for item in items:
        digest = key(item)
        if digest in seen:
            continue
        seen.add(digest)
        result.append(item)
    return result
